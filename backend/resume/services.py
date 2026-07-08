import json
import logging
import re

import google.generativeai as genai
import pdfplumber
from django.conf import settings
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

genai.configure(api_key=settings.GEMINI_API_KEY)
_model = genai.GenerativeModel("gemini-2.5-flash")

ALLOWED_EXTENSIONS = {".pdf", ".docx"}
MAX_UPLOAD_SIZE_BYTES = 5 * 1024 * 1024  # 5 MB
MIN_EXTRACTED_TEXT_LENGTH = 50  # below this, treat as "couldn't read this resume"


class ResumeValidationError(Exception):
    """Raised for problems the candidate can fix (wrong file type, too large, empty)."""


class ResumeExtractionError(Exception):
    """Raised for extraction/parsing failures - not necessarily the candidate's fault."""


def validate_resume_file(uploaded_file) -> str:
    """
    Validates file type and size server-side (never trust a frontend
    `accept=` attribute alone - it's trivially bypassed).
    Returns the lowercase file extension on success.
    """
    name = uploaded_file.name or ""
    ext = "." + name.rsplit(".", 1)[-1].lower() if "." in name else ""

    if ext not in ALLOWED_EXTENSIONS:
        raise ResumeValidationError("Only PDF and DOCX files are supported.")

    if uploaded_file.size > MAX_UPLOAD_SIZE_BYTES:
        raise ResumeValidationError("File is too large. Maximum size is 5MB.")

    if uploaded_file.size == 0:
        raise ResumeValidationError("The uploaded file is empty.")

    return ext


def extract_text_from_pdf(file_obj) -> str:
    text_parts = []

    with pdfplumber.open(file_obj) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

    return "\n".join(text_parts).strip()


def extract_text_from_docx(file_obj) -> str:
    document = DocxDocument(file_obj)
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    # Tables (skills/experience are often laid out in tables in resume templates)
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return "\n".join(paragraphs).strip()


def extract_text(uploaded_file, ext: str) -> str:
    uploaded_file.seek(0)

    try:
        if ext == ".pdf":
            text = extract_text_from_pdf(uploaded_file)
        else:
            text = extract_text_from_docx(uploaded_file)
    except Exception as exc:
        raise ResumeExtractionError(f"Couldn't read this file: {exc}") from exc

    if len(text) < MIN_EXTRACTED_TEXT_LENGTH:
        raise ResumeExtractionError(
            "Couldn't extract readable text from this file. "
            "It may be a scanned image rather than a text-based document."
        )

    return text


def _extract_json_object(raw_text: str) -> dict:
    text = raw_text.strip()

    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    text = re.sub(r"^```(json)?", "", text.strip())
    text = re.sub(r"```$", "", text.strip())

    try:
        return json.loads(text.strip())
    except json.JSONDecodeError:
        pass

    start = text.find("{")
    end = text.rfind("}")

    if start == -1 or end == -1 or end <= start:
        raise ResumeExtractionError("AI response did not contain a parseable JSON object.")

    return json.loads(text[start : end + 1])


_EMPTY_PARSED_DATA = {
    "name": None,
    "email": None,
    "phone": None,
    "education": [],
    "skills": [],
    "experience": [],
    "projects": [],
    "certifications": [],
    "cgpa": None,
    "languages": [],
    "frameworks": [],
}


def structure_resume_with_ai(raw_text: str) -> dict:
    """
    Uses Gemini to turn raw resume text into structured fields. Falls
    back to a best-effort heuristic extraction if the AI call fails,
    so an upload never just dead-ends with nothing to show.
    """
    prompt = f"""
Extract structured information from this resume text.

Return ONLY a raw JSON object, no markdown fences, no commentary, in this exact shape:
{{
  "name": "candidate full name or null",
  "email": "email address found in the resume or null",
  "phone": "phone number found in the resume or null",
  "education": [{{"degree": "...", "institution": "...", "year": "..."}}],
  "skills": ["..."],
  "experience": [{{"title": "...", "company": "...", "duration": "...", "description": "..."}}],
  "projects": [{{"name": "...", "description": "..."}}],
  "certifications": ["..."],
  "cgpa": "value found, as a string, or null",
  "languages": ["..."],
  "frameworks": ["..."]
}}

If a field cannot be found, use null (for single values) or an empty array (for lists).

Resume text:
\"\"\"
{raw_text[:12000]}
\"\"\"
"""

    try:
        response = _model.generate_content(prompt)
        parsed = _extract_json_object(response.text)
        merged = {**_EMPTY_PARSED_DATA, **parsed}
        return merged
    except Exception:
        logger.exception("AI resume structuring failed, falling back to heuristic extraction.")
        return _heuristic_extract(raw_text)


def _heuristic_extract(raw_text: str) -> dict:
    """
    Best-effort, non-AI fallback: regex for email/phone, keyword
    matching for common skills. Not as good as the AI version, but
    keeps the feature usable if Gemini is unavailable.
    """
    data = dict(_EMPTY_PARSED_DATA)

    email_match = re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", raw_text)
    if email_match:
        data["email"] = email_match.group(0)

    phone_match = re.search(r"(\+?\d{1,3}[\s-]?)?\d{10}", raw_text)
    if phone_match:
        data["phone"] = phone_match.group(0)

    cgpa_match = re.search(r"\bCGPA[:\s]*([\d.]+)", raw_text, re.IGNORECASE)
    if cgpa_match:
        data["cgpa"] = cgpa_match.group(1)

    common_skills = [
        "Python", "Java", "JavaScript", "TypeScript", "C++", "C", "SQL",
        "React", "Next.js", "Node.js", "Django", "Flask", "Express",
        "HTML", "CSS", "Tailwind", "MongoDB", "PostgreSQL", "MySQL",
        "Git", "Docker", "Kubernetes", "AWS", "Machine Learning",
        "Data Structures", "Algorithms", "REST API", "GraphQL",
    ]
    found_skills = [s for s in common_skills if re.search(rf"\b{re.escape(s)}\b", raw_text, re.IGNORECASE)]
    data["skills"] = found_skills
    data["frameworks"] = [s for s in found_skills if s in (
        "React", "Next.js", "Django", "Flask", "Express", "Node.js"
    )]

    common_languages = ["English", "Hindi", "Spanish", "French", "German", "Mandarin"]
    data["languages"] = [l for l in common_languages if re.search(rf"\b{l}\b", raw_text, re.IGNORECASE)]

    return data
